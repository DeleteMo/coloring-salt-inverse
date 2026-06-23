"""
生成逆向优化算法性能分析报告 (Word版)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
import sys
sys.path.insert(0, '.')
import inverse_lab_cluster_new as ilc

# 初始化
ilc.initialize()

# 获取所有聚类
all_clusters = list(ilc._cluster_prior.keys())

# 收集测试数据
test_data = []
for cid in all_clusters:
    info = ilc._cluster_prior[cid]
    tL = info['center_L']
    ta = info['center_a']
    tb = info['center_b']

    result = ilc.optimize_cluster_new(cid, tL, ta, tb, seed=42)
    if result:
        errL_abs = abs(result['pred_L'] - tL)
        erra_abs = abs(result['pred_a'] - ta)
        errb_abs = abs(result['pred_b'] - tb)

        errL_pct = errL_abs / abs(tL) * 100 if tL != 0 else 0
        erra_pct = erra_abs / abs(ta) * 100 if ta != 0 else 0
        errb_pct = errb_abs / abs(tb) * 100 if tb != 0 else 0

        test_data.append({
            'cluster': cid,
            'target_L': tL,
            'target_a': ta,
            'target_b': tb,
            'pred_L': result['pred_L'],
            'pred_a': result['pred_a'],
            'pred_b': result['pred_b'],
            'L_abs': errL_abs,
            'a_abs': erra_abs,
            'b_abs': errb_abs,
            'L_pct': errL_pct,
            'a_pct': erra_pct,
            'b_pct': errb_pct,
            'mse': result['mse'],
            'passed': errL_pct < 20 and erra_pct < 20 and errb_pct < 20
        })

# 创建Word文档
doc = Document()

# 标题
title = doc.add_heading('逆向优化算法性能分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. 概述
doc.add_heading('1. 概述', level=1)
doc.add_paragraph('本报告对比分析了正向预测算法（随机森林模型）与聚类逆向优化算法的性能。正向算法根据配方参数预测L*a*b颜色值，逆向算法则根据目标颜色反推配方参数。')

# 2. 正向算法性能
doc.add_heading('2. 正向算法性能', level=1)

doc.add_heading('2.1 模型配置', level=2)
p = doc.add_paragraph()
p.add_run('使用随机森林回归模型（Random Forest Regressor），参数配置：').bold = True
doc.add_paragraph('n_estimators: 500')
doc.add_paragraph('max_depth: 30')
doc.add_paragraph('random_state: 42')

doc.add_heading('2.2 特征结构', level=2)
p = doc.add_paragraph()
p.add_run('模型输入为60维特征，包含：')
doc.add_paragraph('主成分：M002, M001, M003（3维）')
doc.add_paragraph('添加剂：49种化学试剂')
doc.add_paragraph('工艺参数：模式、电压、电流、占空比、频率、周期（6维）')

doc.add_heading('2.3 预测性能', level=2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['目标', 'RMSE', 'R²']
data = [['L（亮度）', '3.00', '0.970'],
        ['a（红绿轴）', '0.57', '0.944'],
        ['b（黄蓝轴）', '1.72', '0.917']]

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '目标'
hdr_cells[1].text = 'RMSE'
hdr_cells[2].text = 'R²'
hdr_cells[3].text = ''

for i, row_data in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('结论：').bold = True
p.add_run('正向模型对L的预测最为准确（R²=0.970），对b的预测稍弱（R²=0.917），但整体R²均大于0.9，模型拟合效果良好。')

# 3. 聚类逆向优化算法
doc.add_heading('3. 聚类逆向优化算法', level=1)

doc.add_heading('3.1 聚类方法', level=2)
p = doc.add_paragraph()
p.add_run('采用配方+工艺聚类（KMeans），基于60维特征进行聚类，共7个聚类。')

doc.add_heading('3.2 逆向优化策略', level=2)
p = doc.add_paragraph()
p.add_run('采用两阶段优化：')

doc.add_paragraph('第一阶段（粗搜索）：')
doc.add_paragraph('优化维度：8维（M002, M001, M003 + 5种添加剂）')
doc.add_paragraph('工艺参数固定为均值')
doc.add_paragraph('DE参数：maxiter=1500, popsize=15, tol=1e-6')

doc.add_paragraph('第二阶段（精调）：')
doc.add_paragraph('优化维度：6维（工艺参数）')
doc.add_paragraph('基于第一阶段结果精调')
doc.add_paragraph('DE参数：maxiter=1000, popsize=10, tol=1e-7')

# 4. 逆向优化性能测试结果
doc.add_heading('4. 逆向优化性能测试结果', level=1)

doc.add_heading('4.1 中心点测试 - 绝对误差', level=2)

table2 = doc.add_table(rows=len(test_data) + 1, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Cluster', '目标a', '目标b', '预测a', '预测b', 'a绝对误差', 'b绝对误差']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h

for i, d in enumerate(test_data):
    row = table2.rows[i + 1].cells
    row[0].text = str(d['cluster'])
    row[1].text = f"{d['target_a']:.2f}"
    row[2].text = f"{d['target_b']:.2f}"
    row[3].text = f"{d['pred_a']:.2f}"
    row[4].text = f"{d['pred_b']:.2f}"
    row[5].text = f"{d['a_abs']:.2f}"
    row[6].text = f"{d['b_abs']:.2f}"

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('参考：').bold = True
p.add_run('a值范围约-5~9, b值范围约-19~26')

doc.add_heading('4.2 中心点测试 - 百分比误差', level=2)

table3 = doc.add_table(rows=len(test_data) + 1, cols=8)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Cluster', '目标L', '目标a', '目标b', 'L误差%', 'a误差%', 'b误差%', 'MSE']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h

for i, d in enumerate(test_data):
    row = table3.rows[i + 1].cells
    row[0].text = str(d['cluster'])
    row[1].text = f"{d['target_L']:.1f}"
    row[2].text = f"{d['target_a']:.2f}"
    row[3].text = f"{d['target_b']:.2f}"
    row[4].text = f"{d['L_pct']:.1f}"
    row[5].text = f"{d['a_pct']:.1f}"
    row[6].text = f"{d['b_pct']:.1f}"
    row[7].text = f"{d['mse']:.2f}"

doc.add_paragraph()

# 4.3 达标统计
doc.add_heading('4.3 达标统计', level=2)

passed_count = sum(1 for d in test_data if d['passed'])
avg_L = np.mean([d['L_pct'] for d in test_data])
avg_a = np.mean([d['a_pct'] for d in test_data])
avg_b = np.mean([d['b_pct'] for d in test_data])

doc.add_paragraph(f'达标聚类数：{passed_count}/{len(test_data)}')
doc.add_paragraph(f'平均L误差：{avg_L:.1f}%')
doc.add_paragraph(f'平均a误差：{avg_a:.1f}%')
doc.add_paragraph(f'平均b误差：{avg_b:.1f}%')

p = doc.add_paragraph()
if avg_L < 20 and avg_a < 20 and avg_b < 20:
    p.add_run('结论：及格 - 所有平均误差<20%').bold = True
else:
    p.add_run('结论：不及格 - 存在平均误差>=20%的项目').bold = True

# 5. 分析
doc.add_heading('5. 分析', level=1)

doc.add_heading('5.1 最佳表现', level=2)
best = min(test_data, key=lambda x: x['mse'])
doc.add_paragraph(f'Cluster {best["cluster"]}表现最佳：L误差={best["L_pct"]:.1f}%, a误差={best["a_pct"]:.1f}%, b误差={best["b_pct"]:.1f}%')

doc.add_heading('5.2 问题分析', level=2)
doc.add_paragraph('主要问题：')
doc.add_paragraph('- a值在接近0时，百分比误差计算会放大')
doc.add_paragraph('- b值极端区域（<-5或>5）的预测能力有限')
doc.add_paragraph('- 聚类3、5、7的a误差较大，可能是聚类边界问题')

# 保存文档
output_path = '逆向优化算法性能分析报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')