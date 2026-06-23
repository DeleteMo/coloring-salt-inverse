"""
生成数据增强方案报告 (Word版) - a值和b值完整分析
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('数据增强方案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 问题
doc.add_heading('问题', level=1)
doc.add_paragraph('逆向优化达标率28.6% (2/7聚类)，原因：a<-2和b<-5区域数据稀疏')

# ==================== b值分析 ====================
doc.add_heading('一、b值分析', level=1)

doc.add_heading('1.1 数据分布', level=2)
table1 = doc.add_table(rows=3, cols=3)
table1.style = 'Table Grid'
table1.rows[0].cells[0].text = '区域'
table1.rows[0].cells[1].text = '样本数'
table1.rows[0].cells[2].text = '比例'
table1.rows[1].cells[0].text = 'b < -5'
table1.rows[1].cells[1].text = '463'
table1.rows[1].cells[2].text = '19.7%'
table1.rows[2].cells[0].text = 'b < -10'
table1.rows[2].cells[1].text = '48'
table1.rows[2].cells[2].text = '2.0%'

doc.add_heading('1.2 关键特征与b值相关性', level=2)
table2 = doc.add_table(rows=3, cols=3)
table2.style = 'Table Grid'
table2.rows[0].cells[0].text = '特征'
table2.rows[0].cells[1].text = '相关系数'
table2.rows[0].cells[2].text = '说明'
table2.rows[1].cells[0].text = 'C007'
table2.rows[1].cells[1].text = 'r = -0.437'
table2.rows[1].cells[2].text = '负相关：减少C007可提高b'
table2.rows[2].cells[0].text = 'C061'
table2.rows[2].cells[1].text = 'r = +0.424'
table2.rows[2].cells[2].text = '正相关：增加C061可提高b'

doc.add_heading('1.3 低b区 vs 高b区特征对比', level=2)
table3 = doc.add_table(rows=5, cols=4)
table3.style = 'Table Grid'
table3.rows[0].cells[0].text = '参数'
table3.rows[0].cells[1].text = '低b区均值(b<-5)'
table3.rows[0].cells[2].text = '高b区均值(b>=5)'
table3.rows[0].cells[3].text = '调整方向'
table3.rows[1].cells[0].text = 'C007'
table3.rows[1].cells[1].text = '4.79'
table3.rows[1].cells[2].text = '0.15'
table3.rows[1].cells[3].text = '↓ 减少'
table3.rows[2].cells[0].text = 'C061'
table3.rows[2].cells[1].text = '0.00'
table3.rows[2].cells[2].text = '8.52'
table3.rows[2].cells[3].text = '↑ 增加'
table3.rows[3].cells[0].text = 'M001'
table3.rows[3].cells[1].text = '0.00'
table3.rows[3].cells[2].text = '4.19'
table3.rows[3].cells[3].text = '↑ 增加'
table3.rows[4].cells[0].text = '电流'
table3.rows[4].cells[1].text = '23.04'
table3.rows[4].cells[2].text = '14.93'
table3.rows[4].cells[3].text = '↓ 降低'

doc.add_heading('1.4 配方调整方案', level=2)
doc.add_paragraph('减少：C007(-4.6)、电流(-8A)、占空比(-6%)')
doc.add_paragraph('增加：C061(+8)、M001(+4)')

# ==================== a值分析 ====================
doc.add_heading('二、a值分析', level=1)

doc.add_heading('2.1 数据分布', level=2)
table4 = doc.add_table(rows=2, cols=3)
table4.style = 'Table Grid'
table4.rows[0].cells[0].text = '区域'
table4.rows[0].cells[1].text = '样本数'
table4.rows[0].cells[2].text = '比例'
table4.rows[1].cells[0].text = 'a < -2'
table4.rows[1].cells[1].text = '184'
table4.rows[1].cells[2].text = '7.8%'

doc.add_heading('2.2 关键特征与a值相关性', level=2)
table5 = doc.add_table(rows=3, cols=3)
table5.style = 'Table Grid'
table5.rows[0].cells[0].text = '特征'
table5.rows[0].cells[1].text = '相关系数'
table5.rows[0].cells[2].text = '说明'
table5.rows[1].cells[0].text = 'M002'
table5.rows[1].cells[1].text = 'r = +0.312'
table5.rows[1].cells[2].text = '正相关：高M002导致低a'
table5.rows[2].cells[0].text = '电流'
table5.rows[2].cells[1].text = 'r = -0.289'
table5.rows[2].cells[2].text = '负相关：高电流导致低a'

doc.add_heading('2.3 低a区 vs 正常a区特征对比', level=2)
table6 = doc.add_table(rows=5, cols=4)
table6.style = 'Table Grid'
table6.rows[0].cells[0].text = '参数'
table6.rows[0].cells[1].text = '低a区均值(a<-2)'
table6.rows[0].cells[2].text = '正常a区均值(a>=-2)'
table6.rows[0].cells[3].text = '调整方向'
table6.rows[1].cells[0].text = 'M002'
table6.rows[1].cells[1].text = '37.80'
table6.rows[1].cells[2].text = '28.37'
table6.rows[1].cells[3].text = '↓ 减少'
table6.rows[2].cells[0].text = 'M001'
table6.rows[2].cells[1].text = '0.16'
table6.rows[2].cells[2].text = '5.35'
table6.rows[2].cells[3].text = '↑ 增加'
table6.rows[3].cells[0].text = '电流'
table6.rows[3].cells[1].text = '20.31'
table6.rows[3].cells[2].text = '12.70'
table6.rows[3].cells[3].text = '↓ 降低'
table6.rows[4].cells[0].text = 'C007'
table6.rows[4].cells[1].text = '5.72'
table6.rows[4].cells[2].text = '2.07'
table6.rows[4].cells[3].text = '↓ 减少'

doc.add_heading('2.4 配方调整方案', level=2)
doc.add_paragraph('减少：M002(-9)、电流(-8A)、C007(-3.7)')
doc.add_paragraph('增加：M001(+5)、C013(+0.8)')

# ==================== 综合方案 ====================
doc.add_heading('三、综合实施方案', level=1)

table7 = doc.add_table(rows=3, cols=4)
table7.style = 'Table Grid'
table7.rows[0].cells[0].text = '优先级'
table7.rows[0].cells[1].text = '目标'
table7.rows[0].cells[2].text = '新增样本'
table7.rows[0].cells[3].text = '核心配方'
table7.rows[1].cells[0].text = '高'
table7.rows[1].cells[1].text = 'a < -2'
table7.rows[1].cells[2].text = '80-100'
table7.rows[1].cells[3].text = 'M002↓、电流↓、M001↑'
table7.rows[2].cells[0].text = '高'
table7.rows[2].cells[1].text = 'b < -10'
table7.rows[2].cells[2].text = '100-150'
table7.rows[2].cells[3].text = 'C007↓、C061↑、电流↓'

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('预期效果：达标率 28.6% → 60%+').bold = True

# ==================== 可视化 ====================
doc.add_heading('四、可视化分析', level=1)

try:
    doc.add_picture('correlation_analysis.png', width=Inches(5))
    doc.add_paragraph('图1：特征与b值相关性分析')
except:
    pass

try:
    doc.add_picture('a_value_analysis.png', width=Inches(5))
    doc.add_paragraph('图2：特征与a值相关性分析')
except:
    pass

try:
    doc.add_picture('distribution_analysis.png', width=Inches(5))
    doc.add_paragraph('图3：a/b值分布及特征对比')
except:
    pass

doc.save('data_enhancement_report_final.docx')
print('已生成: 数据增强方案_完整版.docx')